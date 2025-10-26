"""
S3 Document Loader
Loads documents from S3 bucket and extracts text
"""
import os
import io
import boto3
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️  python-docx not installed. Install with: pip install python-docx")


class S3DocumentLoader:
    def __init__(
        self,
        bucket_name: str,
        prefix: str = "",
        aws_access_key_id: str = None,
        aws_secret_access_key: str = None,
        region_name: str = "us-east-1"
    ):
        self.bucket_name = bucket_name
        self.prefix = prefix

        # Initialize S3 client
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=aws_access_key_id or os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=aws_secret_access_key or os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=region_name or os.getenv('AWS_REGION', 'us-east-1')
        )

        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a .docx file"""
        if not HAS_DOCX:
            return ""

        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            print(f"  ⚠️  Error parsing .docx: {e}")
            return ""

    def load_documents(self) -> List[Document]:
        """Load all documents from S3 bucket using fast custom loader"""
        print(f"\n🔍 Searching S3 path: s3://{self.bucket_name}/{self.prefix}")

        # First, list what files are available
        files = self.list_documents()
        if not files:
            print(f"⚠️  No files found in s3://{self.bucket_name}/{self.prefix}")
            return []

        print(f"✓ Found {len(files)} files:")
        for f in files[:10]:  # Show first 10 files
            print(f"  - {f}")
        if len(files) > 10:
            print(f"  ... and {len(files) - 10} more")

        # Load documents using fast custom method
        documents = []
        print(f"⏳ Loading {len(files)} documents...")

        for i, key in enumerate(files, 1):
            try:
                print(f"  [{i}/{len(files)}] Loading {key.split('/')[-1]}...")

                # Download file from S3
                response = self.s3_client.get_object(Bucket=self.bucket_name, Key=key)
                file_content = response['Body'].read()

                # Extract text based on file type
                if key.endswith('.docx'):
                    text = self._extract_text_from_docx(file_content)
                elif key.endswith('.txt'):
                    text = file_content.decode('utf-8')
                else:
                    print(f"  ⚠️  Unsupported file type: {key}")
                    continue

                if text:
                    doc = Document(
                        page_content=text,
                        metadata={"source": key}
                    )
                    documents.append(doc)
                    print(f"  ✓ Loaded {len(text)} characters")

            except Exception as e:
                print(f"  ❌ Error loading {key}: {e}")
                continue

        print(f"✓ Successfully loaded {len(documents)} documents from S3")
        return documents

    def load_and_split(self) -> List[Document]:
        """Load documents and split into chunks"""
        documents = self.load_documents()

        if not documents:
            return []

        # Split documents into chunks
        chunks = self.text_splitter.split_documents(documents)
        print(f"Split into {len(chunks)} chunks")

        return chunks

    def list_documents(self) -> List[str]:
        """List all document keys in the S3 bucket"""
        try:
            response = self.s3_client.list_objects_v2(
                Bucket=self.bucket_name,
                Prefix=self.prefix
            )

            if 'Contents' not in response:
                return []

            keys = [obj['Key'] for obj in response['Contents']]
            print(f"Found {len(keys)} files in S3")
            return keys

        except Exception as e:
            print(f"Error listing S3 objects: {e}")
            return []


if __name__ == "__main__":
    # Test the loader
    from dotenv import load_dotenv
    load_dotenv()

    loader = S3DocumentLoader(
        bucket_name=os.getenv('S3_BUCKET_NAME'),
        prefix=os.getenv('S3_DOCUMENTS_PREFIX', '')
    )

    # List files
    files = loader.list_documents()
    print(f"Files: {files}")

    # Load and split
    chunks = loader.load_and_split()
    print(f"Loaded {len(chunks)} chunks")
    if chunks:
        print(f"Sample chunk: {chunks[0].page_content[:200]}")
